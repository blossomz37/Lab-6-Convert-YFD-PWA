#!/usr/bin/env python3
"""
Convert markdown to DOCX with exact ProWriting Aid compatible structure
Matches the structure of document.docx (Heading1 + Normal styles only)
"""

import sys
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("Error: python-docx not installed. Install with: pip install python-docx")
    sys.exit(1)

def parse_markdown_content(markdown_text):
    """Parse markdown content and extract chapters with their content"""
    chapters = []
    current_chapter = None
    
    lines = markdown_text.split('\n')
    
    for line in lines:
        # Check for chapter heading
        chapter_match = re.match(r'^# Chapter (\d+): (.+)$', line.strip())
        if chapter_match:
            # Save previous chapter if exists
            if current_chapter:
                chapters.append(current_chapter)
            
            # Start new chapter
            current_chapter = {
                'number': int(chapter_match.group(1)),
                'title': chapter_match.group(2).strip(),
                'content_paragraphs': []
            }
        elif line.strip() and current_chapter:
            # Add content paragraph (skip empty lines)
            current_chapter['content_paragraphs'].append(line.strip())
    
    # Add the last chapter
    if current_chapter:
        chapters.append(current_chapter)
    
    return chapters

def create_prowriting_aid_docx(markdown_file, output_file):
    """Create DOCX with exact ProWriting Aid compatible structure"""
    
    # Read markdown content
    try:
        with open(markdown_file, 'r', encoding='utf-8') as f:
            markdown_content = f.read()
    except Exception as e:
        print(f"Error reading {markdown_file}: {e}")
        return False
    
    # Parse chapters
    chapters = parse_markdown_content(markdown_content)
    
    if not chapters:
        print("No chapters found in markdown file")
        return False
    
    # Create new document
    doc = Document()
    
    # Clear default styles and ensure we have the right ones
    styles = doc.styles
    
    # Ensure Heading1 style exists and is configured
    try:
        heading1_style = styles['Heading 1']
    except KeyError:
        heading1_style = styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
    
    # Ensure Normal style exists
    try:
        normal_style = styles['Normal']
    except KeyError:
        normal_style = styles.add_style('Normal', WD_STYLE_TYPE.PARAGRAPH)
    
    print(f"Creating DOCX with {len(chapters)} chapters...")
    
    # Process each chapter
    for i, chapter in enumerate(chapters):
        chapter_title = f"Chapter {chapter['number']}: {chapter['title']}"
        
        # Add chapter heading with Heading1 style
        heading_para = doc.add_paragraph(chapter_title)
        heading_para.style = 'Heading 1'
        
        # Add CRITICAL empty paragraph with Normal style (matches document.docx)
        empty_para = doc.add_paragraph()
        empty_para.style = 'Normal'
        
        # Add content paragraphs with Normal style
        for content_para in chapter['content_paragraphs']:
            if content_para.strip():  # Skip empty paragraphs
                para = doc.add_paragraph(content_para)
                para.style = 'Normal'
        
        # Add spacing between chapters (except last chapter)
        if i < len(chapters) - 1:
            # Add extra spacing between chapters
            for _ in range(3):
                spacer = doc.add_paragraph()
                spacer.style = 'Normal'
        
        print(f"  Added: {chapter_title}")
    
    # Save document
    try:
        doc.save(output_file)
        print(f"✅ Successfully created: {output_file}")
        return True
    except Exception as e:
        print(f"Error saving {output_file}: {e}")
        return False

def verify_docx_structure(docx_file):
    """Verify the created DOCX has the correct structure"""
    try:
        doc = Document(docx_file)
        
        styles_used = {}
        chapter_count = 0
        empty_para_count = 0
        
        for para in doc.paragraphs[:20]:  # Check first 20 paragraphs
            style_name = para.style.name if para.style else 'Unknown'
            styles_used[style_name] = styles_used.get(style_name, 0) + 1
            
            if 'Chapter' in para.text and 'Heading' in style_name:
                chapter_count += 1
            elif not para.text.strip() and style_name == 'Normal':
                empty_para_count += 1
        
        print(f"\n📊 Structure Verification:")
        print(f"  Styles used: {dict(styles_used)}")
        print(f"  Chapters found: {chapter_count}")
        print(f"  Empty Normal paragraphs: {empty_para_count}")
        
        # Check if structure matches expectations
        expected_styles = {'Heading 1', 'Normal'}
        actual_styles = set(styles_used.keys())
        
        if actual_styles <= expected_styles:
            print("✅ Structure matches ProWriting Aid requirements!")
            return True
        else:
            print(f"⚠️  Unexpected styles found: {actual_styles - expected_styles}")
            return False
            
    except Exception as e:
        print(f"Error verifying structure: {e}")
        return False

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python markdown_to_docx.py <input_markdown> <output_docx>")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    if not Path(input_file).exists():
        print(f"Error: Input file {input_file} not found")
        sys.exit(1)
    
    # Convert markdown to DOCX
    success = create_prowriting_aid_docx(input_file, output_file)
    
    if success:
        # Verify the structure
        verify_docx_structure(output_file)
    else:
        sys.exit(1)